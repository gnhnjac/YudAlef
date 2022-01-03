import docx
from docx.shared import Cm, Pt
import bs4
import os.path

html_table = input("HTML Table:\n")

soup = bs4.BeautifulSoup(html_table)

table_list = [[]]

tr_iterator = iter(soup.findAll('tr'))

headers = next(tr_iterator)

for cell in headers:

    try:
        table_list[0].append(cell.get_text())
    except:
        continue



for i, next_row in enumerate(tr_iterator):

    table_list.append([])

    for cell in next_row:

        try:
            table_list[i+1].append(cell.get_text())
        except:
            continue

doc = docx.Document()

table = doc.add_table(len(table_list), len(table_list[0]))
table.style = 'TableGrid'

for i in range(len(table_list)):

    for j in range(len(table_list[0])):

        table.cell(i, j).add_paragraph(table_list[i][j])


doc.add_page_break()

doc.save(os.path.dirname(__file__) + '/result.docx')
